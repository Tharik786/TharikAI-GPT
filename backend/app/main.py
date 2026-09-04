import os
import sys
import json
from pathlib import Path
from dotenv import load_dotenv

# Robustly load .env and sys.path from backend directory regardless of working directory
backend_dir = Path(__file__).resolve().parent.parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

load_dotenv(backend_dir / ".env")
load_dotenv()


import io
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.gemini_client import stream_chat_completion, GeminiError
from app.tavily_client import (
    search_tavily,
    format_search_context,
    should_auto_search,
    run_deep_research,
    format_deep_research_context,
)
from app.image_client import detect_image_prompt, generate_image_url
from app.db import (
    init_db,
    get_db_connection,
    get_user,
    save_user,
    list_conversations,
    save_conversation,
    set_messages as db_set_messages,
    delete_conversation as db_delete_conversation,
)


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Initialize cloud PostgreSQL tables on startup
    init_db()
    yield


app = FastAPI(title="TharikAI API", lifespan=lifespan)

cors_env = os.getenv("CORS_ORIGINS", "*").strip()
if cors_env == "*":
    app.add_middleware(
        CORSMiddleware,
        allow_origin_regex=".*",
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    origins = [o.strip() for o in cors_env.split(",") if o.strip()]
    app.add_middleware(
        CORSMiddleware,
        allow_origins=origins,
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )


class ChatMessage(BaseModel):
    role: str  # "user" | "assistant"
    content: str
    images: list[dict | str] | None = None


class ChatBody(BaseModel):
    messages: list[ChatMessage]
    web_search: bool | None = False
    deep_research: bool | None = False
    email: str | None = None


class SearchBody(BaseModel):
    query: str
    max_results: int | None = 5


class ResearchBody(BaseModel):
    query: str


class RegisterBody(BaseModel):
    email: str
    name: str | None = None
    password_hash: str


class LoginBody(BaseModel):
    email: str
    password_hash: str


class ConversationBody(BaseModel):
    id: str
    email: str
    title: str = "New chat"
    createdAt: int
    updatedAt: int


class MessagesBody(BaseModel):
    id: str
    messages: list[dict]
    updatedAt: int



@app.api_route("/", methods=["GET", "HEAD"])
async def root():
    return {"status": "ok", "message": "TharikAI API is running", "docs": "/docs"}


@app.api_route("/api/health", methods=["GET", "HEAD"])
async def health():

    db_status = "disconnected"
    try:
        con = get_db_connection()
        con.run("SELECT 1")
        con.close()
        db_status = "connected"
    except Exception as e:
        db_status = f"error: {str(e)}"

    return {
        "status": "ok" if db_status == "connected" else "degraded",
        "database": db_status,
    }


@app.post("/api/auth/register")
async def register(body: RegisterBody):
    email = body.email.strip().lower()
    existing = get_user(email)
    if existing:
        raise HTTPException(status_code=400, detail="An account with this email already exists.")
    user = save_user(email, body.name or email.split("@")[0], body.password_hash)
    return {"success": True, "user": {"email": user["email"], "name": user["name"]}}


@app.post("/api/auth/login")
async def login(body: LoginBody):
    email = body.email.strip().lower()
    user = get_user(email)
    if not user or user.get("password_hash") != body.password_hash:
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    return {"success": True, "user": {"email": user["email"], "name": user["name"]}}


@app.get("/api/conversations")
async def get_user_conversations(email: str):
    email = email.strip().lower()
    convs = list_conversations(email)
    return {"conversations": convs}


@app.post("/api/conversations")
async def create_or_update_conversation(body: ConversationBody):
    save_conversation(body.id, body.email, body.title, body.createdAt, body.updatedAt)
    return {"success": True}


@app.post("/api/conversations/messages")
async def update_conversation_messages(body: MessagesBody):
    db_set_messages(body.id, body.messages, body.updatedAt)
    return {"success": True}


@app.delete("/api/conversations/{conv_id}")
async def remove_conversation(conv_id: str):
    db_delete_conversation(conv_id)
    return {"success": True}


@app.post("/api/extract-document")
async def extract_document(file: UploadFile = File(...)):
    """
    Extracts clean readable text from uploaded documents (PDF, DOCX, TXT, MD, CSV, code).
    Prevents binary byte corruption and enables AI models to read and analyze files.
    """
    filename = file.filename or "uploaded_file"
    ext = filename.lower().split(".")[-1] if "." in filename else ""
    content = await file.read()
    
    extracted_text = ""
    page_count = 1

    if ext == "pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            page_count = len(reader.pages)
            pages_text = []
            for idx, page in enumerate(reader.pages):
                txt = page.extract_text() or ""
                if txt.strip():
                    pages_text.append(f"--- Page {idx + 1} ---\n{txt.strip()}")
            extracted_text = "\n\n".join(pages_text).strip()
            if not extracted_text:
                extracted_text = "[Notice: This PDF contains no extractable text. It may contain scanned images or protected content.]"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read PDF: {str(e)}")

    elif ext in ("docx", "doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            extracted_text = "\n\n".join(paragraphs).strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read Word document: {str(e)}")

    else:
        # Plain text, Markdown, CSV, JSON, code files
        try:
            extracted_text = content.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = content.decode("latin-1", errors="ignore")

    return {
        "success": True,
        "filename": filename,
        "page_count": page_count,
        "size": len(content),
        "text": extracted_text,
    }


@app.post("/api/research")
async def deep_research_endpoint(body: ResearchBody):
    """
    Dedicated multi-source deep research query endpoint.
    """
    if not body.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    data = await run_deep_research(body.query.strip())
    return data


@app.post("/api/search")
async def search_endpoint(body: SearchBody):
    """
    Direct web search endpoint powered by Tavily Search API.
    """
    if not body.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    data = await search_tavily(body.query.strip(), max_results=body.max_results or 5)
    return data


class ImageBody(BaseModel):
    prompt: str
    width: int | None = 1024
    height: int | None = 1024
    model: str | None = "flux"


@app.post("/api/image")
async def image_endpoint(body: ImageBody):
    """
    Direct AI image generation endpoint powered by Cloudflare Worker Image API.
    """
    if not body.prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt cannot be empty")
    return await generate_image_url(body.prompt.strip())


@app.post("/api/chat")
async def chat(body: ChatBody):

    """
    Unified streaming endpoint:
    - Multimodal vision (image input)
    - Real-time Web Search and Deep Multi-Source Research
    - AI Image Generation (Cloudflare Worker)
    """
    if not body.messages:
        raise HTTPException(status_code=400, detail="messages cannot be empty")

    llm_messages = []
    for m in body.messages:
        msg_dict = {"role": m.role, "content": m.content}
        if m.images:
            msg_dict["images"] = m.images
        llm_messages.append(msg_dict)

    # Find the latest user message
    last_user_msg = next((m.content for m in reversed(body.messages) if m.role == "user"), "")
    clean_search_query = last_user_msg.split("--- Document Attached:")[0].split("[Attached image:")[0].strip()

    # Detect if user requested deep research
    is_deep_research = bool(
        body.deep_research
        or clean_search_query.lower().startswith("/research")
        or clean_search_query.lower().startswith("deep research:")
    )
    if clean_search_query.lower().startswith("/research"):
        clean_search_query = clean_search_query[9:].strip()

    # Detect if user is asking to generate an AI image
    image_prompt = detect_image_prompt(clean_search_query) if not is_deep_research else None

    # Auto web search check
    perform_search = (
        should_auto_search(clean_search_query)
        if (body.web_search is not False and not image_prompt and not is_deep_research)
        else False
    )

    async def event_stream():
        # 1. Handle AI Image Generation Intent
        if image_prompt:
            try:
                yield f"data: {json.dumps({'type': 'search_status', 'status': f'Generating AI image for \"{image_prompt[:40]}\"...'})}\n\n"
                img_data = await generate_image_url(image_prompt)
                intro = f"Here is the generated image of **{image_prompt}**:\n\n"
                for word in intro.split(" "):
                    yield f"data: {json.dumps({'delta': word + ' '})}\n\n"
                img_markdown = f"![{image_prompt}]({img_data['image_url']})\n\n"
                yield f"data: {json.dumps({'delta': img_markdown})}\n\n"
                yield f"data: {json.dumps({'done': True})}\n\n"
                return
            except Exception as e:
                yield f"data: {json.dumps({'error': f'Image generation error: {str(e)}'})}\n\n"
                return


        # 2. Handle Deep Multi-Source Research Intent
        if is_deep_research and clean_search_query:
            try:
                yield f"data: {json.dumps({'type': 'search_status', 'status': f'🔍 Deep Research: Formulating multi-source sub-queries for \"{clean_search_query[:45]}\"...'})}\n\n"
                research_data = await run_deep_research(clean_search_query)
                if research_data.get("success") and research_data.get("results"):
                    source_count = len(research_data.get("results", []))
                    status_text = f"🌐 Deep Research: Aggregated {source_count} verified sources. Synthesizing comprehensive report..."
                    yield f"data: {json.dumps({'type': 'search_status', 'status': status_text})}\n\n"
                    yield f"data: {json.dumps({'type': 'sources', 'sources': research_data['results']})}\n\n"
                    deep_context = format_deep_research_context(research_data)
                    
                    async for chunk in stream_chat_completion(
                        llm_messages,
                        web_search_context=deep_context,
                    ):
                        yield f"data: {json.dumps({'delta': chunk})}\n\n"
                    yield f"data: {json.dumps({'done': True})}\n\n"
                    return
            except Exception as e:
                print(f"Deep research fallback note: {e}")

        # 3. Handle Standard Live Web Search Intent
        search_context = ""
        if perform_search and clean_search_query:
            try:
                yield f"data: {json.dumps({'type': 'search_status', 'status': f'Searching the web for \"{clean_search_query[:50]}\"...'})}\n\n"
                search_data = await search_tavily(clean_search_query, max_results=5)
                if search_data.get("success") and search_data.get("results"):
                    yield f"data: {json.dumps({'type': 'sources', 'sources': search_data['results']})}\n\n"
                    search_context = format_search_context(search_data)
                elif search_data.get("error"):
                    yield f"data: {json.dumps({'type': 'search_status', 'status': ''})}\n\n"
            except Exception as e:
                print(f"Web search non-blocking error: {e}")

        # 4. Stream LLM Chat Completion (Multimodal Vision + Grounding)
        try:
            async for chunk in stream_chat_completion(
                llm_messages,
                web_search_context=search_context,
            ):
                yield f"data: {json.dumps({'delta': chunk})}\n\n"
        except GeminiError as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            return
        yield f"data: {json.dumps({'done': True})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "8000"))
    uvicorn.run("app.main:app", host="0.0.0.0", port=port, reload=True)

